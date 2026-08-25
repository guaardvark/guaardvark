#!/usr/bin/env python3
"""Generate the RAG validation corpus.

Four documents, each shaped to exercise a failure mode this pipeline has
actually produced, rather than a hypothetical one:

  flat_prose_no_headings.md   Large prose with no headings at all. Sectioning
                              keys off headings, so a document without any is
                              the case that once produced a 105x chunk
                              explosion.
  structured_manual.md        Deep heading hierarchy, to check that the
                              heading path survives into chunk metadata and
                              that sections chunk at their own boundaries.
  mixed_prose_and_code.md     Alternating prose and fenced code, to check
                              prose is not chunked at code scale.
  field_report.docx           A real binary, so the Docling path runs and
                              page provenance can be checked.

Content is generated from a fixed seed: the corpus is reproducible from this
script, so the repository carries a generator rather than megabytes of
synthetic prose. Each document hides verifiable facts ("needles") that
questions.json asks about, so retrieval can be scored objectively instead of
by reading the output and forming an impression.
"""

import argparse
import json
import pathlib
import random

SEED = 20260825

# Deliberately mundane vocabulary. Distinctive wording would let a retriever
# match on style rather than meaning, which is not what is being measured.
WORDS = """station module coolant valve rotation sequence operator manifest
telemetry pressure vessel intake exhaust filter cycle calibration drift
tolerance reading sensor array conduit relay switch panel gauge threshold
baseline nominal deviation inspection interval maintenance schedule roster
shift handover checklist procedure clearance authorisation log entry record
sample analysis result variance margin corrective action report summary""".split()

# The facts each document hides. Distinctive enough to score exactly, phrased
# plainly enough that a retriever has to find the passage rather than pattern
# match on an odd token.
NEEDLES = {
    "flat": "The Meridian Protocol was ratified on 14 March 2019 by delegates from twelve member stations.",
    "structured": "The emergency purge valve must be rotated counterclockwise exactly three times before venting.",
    "code_prose": "Throughput collapsed once the retry budget exceeded four attempts per request.",
    "code_body": "MAX_RETRY_BUDGET = 4  # above this, throughput collapses",
    "docx": "Sample K-19 recorded a coolant drift of 0.42 units per hour, the highest in the survey.",
    "tiny": "Relay bypass authority for Module 7 rests with the duty supervisor, not the shift lead.",
    # Deliberately placed in two documents: a correct ranking puts the fuller
    # treatment first, which a single-document needle cannot measure.
    "shared": "Calibration drift above 0.30 units per hour requires a same-shift inspection.",
}


def para(rng, n_sentences=6):
    out = []
    for _ in range(n_sentences):
        n = rng.randint(9, 18)
        s = " ".join(rng.choice(WORDS) for _ in range(n))
        out.append(s[0].upper() + s[1:] + ".")
    return " ".join(out)


def make_tiny(rng):
    """A short note -- the shape most of a real corpus is made of.

    Its value to the benchmark is what it does NOT contain. At a couple of
    kilobytes the per-chunk work is negligible, so whatever this document costs
    is very nearly the fixed price of ingesting any document at all. A corpus of
    large files hides that term completely, and it is the term that decides
    whether a hundred thousand small files is an afternoon or a fortnight.
    """
    return "\n\n".join([
        "# Module 7 Handover Note",
        para(rng, 3),
        NEEDLES["tiny"],
        para(rng, 2),
    ]) + "\n"


def make_flat(rng, target_kb):
    """No headings anywhere -- the sectioning path's degenerate case."""
    parts = [
        "This record is maintained as a continuous narrative. It carries no "
        "headings, sections, or numbered divisions of any kind, by design.",
    ]
    size = 0
    i = 0
    while size < target_kb * 1024:
        parts.append(para(rng))
        size += len(parts[-1])
        i += 1
        if i == 40:
            parts.append(NEEDLES["flat"])
        if i == 80:
            parts.append(NEEDLES["shared"])
    if i < 40:
        parts.append(NEEDLES["flat"])
    return "\n\n".join(parts) + "\n"


def make_structured(rng, target_kb):
    """Four heading levels, so the heading path has something to carry."""
    out = ["# Station Operations Manual", "", para(rng), ""]
    size = 0
    placed = False
    for chapter in range(1, 7):
        out += [f"## {chapter}. Chapter {chapter}", "", para(rng), ""]
        for section in range(1, 5):
            out += [f"### {chapter}.{section} Section {chapter}.{section}", "", para(rng), ""]
            for sub in range(1, 4):
                out += [f"#### {chapter}.{section}.{sub} Procedure {chapter}.{section}.{sub}", ""]
                body = para(rng)
                # Bury the needle deep, where only a surviving heading path
                # makes the passage locatable by its context.
                if chapter == 4 and section == 2 and sub == 3 and not placed:
                    body += " " + NEEDLES["structured"]
                    placed = True
                if chapter == 2 and section == 1 and sub == 1:
                    body += " " + NEEDLES["shared"]
                out += [body, ""]
                size += len(body)
        if size > target_kb * 1024:
            break
    if not placed:
        out += ["#### Appendix Procedure", "", NEEDLES["structured"], ""]
    return "\n".join(out) + "\n"


def make_mixed(rng, target_kb):
    """Prose and code alternating, at code's natural block size."""
    out = ["# Service Notes", ""]
    size = 0
    placed = False
    i = 0
    while size < target_kb * 1024:
        i += 1
        body = para(rng)
        if i == 12 and not placed:
            body += " " + NEEDLES["code_prose"]
            placed = True
        out += [body, ""]
        code = [
            "```python",
            f"def handler_{i}(request, retries=0):",
            '    """Retry with a bounded budget."""',
        ]
        if i == 12:
            code.append("    " + NEEDLES["code_body"])
        code += [
            "    while retries < 4:",
            "        result = dispatch(request)",
            "        if result.ok:",
            "            return result",
            "        retries += 1",
            "    raise TimeoutError('budget exhausted')",
            "```",
        ]
        out += code + [""]
        size += len(body) + sum(len(c) for c in code)
    if not placed:
        out += [NEEDLES["code_prose"], ""]
    return "\n".join(out) + "\n"


def make_docx(rng, path, pages=6):
    from docx import Document as Docx
    d = Docx()
    d.add_heading("Coolant Survey Field Report", level=1)
    for p in range(1, pages + 1):
        d.add_heading(f"Site {p}", level=2)
        for _ in range(4):
            d.add_paragraph(para(rng))
        # Late enough that finding it proves later pages were parsed, not just
        # the first.
        if p == 4:
            d.add_paragraph(NEEDLES["docx"])
        if p < pages:
            d.add_page_break()
    d.save(str(path))


def main():
    ap = argparse.ArgumentParser(description=__doc__,
                                 formatter_class=argparse.RawDescriptionHelpFormatter)
    ap.add_argument("--out", required=True, help="Directory to write the corpus into")
    ap.add_argument("--flat-kb", type=int, default=150,
                    help="Size of the no-headings document (default 150)")
    args = ap.parse_args()

    out = pathlib.Path(args.out)
    out.mkdir(parents=True, exist_ok=True)
    rng = random.Random(SEED)

    written = []
    for name, text in (
        ("tiny_note.md", make_tiny(rng)),
        ("flat_prose_no_headings.md", make_flat(rng, args.flat_kb)),
        ("structured_manual.md", make_structured(rng, 60)),
        ("mixed_prose_and_code.md", make_mixed(rng, 40)),
    ):
        (out / name).write_text(text, encoding="utf-8")
        written.append((name, len(text)))

    try:
        make_docx(rng, out / "field_report.docx")
        written.append(("field_report.docx", (out / "field_report.docx").stat().st_size))
    except ImportError:
        print("python-docx not installed — skipping the Docling fixture")

    (out / "needles.json").write_text(json.dumps(NEEDLES, indent=2), encoding="utf-8")
    for name, size in written:
        print(f"  {name:<32} {size/1024:8.1f} KB")
    print(f"{len(written)} document(s) written to {out}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
